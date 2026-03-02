from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from django.core.mail import send_mail
from django.conf import settings
import io
import base64
from docx import Document
from docx.shared import Inches
from .models import (
    User, Company, CompanyContact, Contract, WorkCenter, 
    Project, ProjectDocument, Meeting, DocumentTemplate, MeetingDocument,
    FollowUpMeeting, FollowUpMeetingConfig
)
from .serializers import (
    UserSerializer, CompanySerializer, CompanyContactSerializer, ContractSerializer, 
    WorkCenterSerializer, ProjectSerializer, ProjectDocumentSerializer, 
    MeetingSerializer, DocumentTemplateSerializer, MeetingDocumentSerializer,
    FollowUpMeetingSerializer, FollowUpMeetingConfigSerializer
)


# ── Shared DOCX template helpers ──

def replace_in_paragraph(para, replacements_dict):
    """
    Replaces placeholder keys in a paragraph while preserving run-level formatting.
    """
    full_text = para.text
    if not any(key in full_text for key in replacements_dict):
        return

    runs = list(para.runs)
    if not runs:
        return

    run_spans = []
    pos = 0
    for run in runs:
        run_spans.append((pos, pos + len(run.text), run))
        pos += len(run.text)

    for key, val in replacements_dict.items():
        search_start = 0
        while True:
            idx = full_text.find(key, search_start)
            if idx == -1:
                break
            key_end = idx + len(key)

            overlapping = [
                (s, e, r) for (s, e, r) in run_spans
                if s < key_end and e > idx
            ]

            if overlapping:
                first_s, first_e, first_run = overlapping[0]

                if len(overlapping) > 1:
                    merged = "".join(r.text for (_, _, r) in overlapping)
                    first_run.text = merged
                    for (_, _, r) in overlapping[1:]:
                        r._element.getparent().remove(r._element)

                first_run.text = first_run.text.replace(key, str(val))

                run_spans = []
                pos2 = 0
                for r in para.runs:
                    run_spans.append((pos2, pos2 + len(r.text), r))
                    pos2 += len(r.text)
                full_text = para.text

            search_start = idx + 1


def replace_in_blocks(blocks, replacements_dict):
    for p in blocks.paragraphs:
        replace_in_paragraph(p, replacements_dict)
    for table in blocks.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements_dict)


class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer

class CompanyViewSet(viewsets.ModelViewSet):
    queryset = Company.objects.all()
    serializer_class = CompanySerializer

class CompanyContactViewSet(viewsets.ModelViewSet):
    queryset = CompanyContact.objects.all()
    serializer_class = CompanyContactSerializer

class ContractViewSet(viewsets.ModelViewSet):
    queryset = Contract.objects.all()
    serializer_class = ContractSerializer

class WorkCenterViewSet(viewsets.ModelViewSet):
    queryset = WorkCenter.objects.all()
    serializer_class = WorkCenterSerializer

class ProjectViewSet(viewsets.ModelViewSet):
    queryset = Project.objects.all()
    serializer_class = ProjectSerializer

class ProjectDocumentViewSet(viewsets.ModelViewSet):
    queryset = ProjectDocument.objects.all()
    serializer_class = ProjectDocumentSerializer

class MeetingViewSet(viewsets.ModelViewSet):
    queryset = Meeting.objects.all()
    serializer_class = MeetingSerializer

    def perform_create(self, serializer):
        meeting = serializer.save()
        try:
            self._generate_meeting_document(meeting)
        except Exception as e:
            print(f"Error generating meeting document: {e}")

    def _generate_meeting_document(self, meeting):
        template = DocumentTemplate.objects.filter(name__icontains="Acta inicial", category="Acta Inicial").first()
        if not template:
            template = DocumentTemplate.objects.filter(category="Acta Inicial").first()

        if template and template.file_data:
            file_data = template.file_data
            if "," in file_data:
                base64_data = file_data.split(",")[1]
            else:
                base64_data = file_data

            doc_bytes = base64.b64decode(base64_data)
            doc = Document(io.BytesIO(doc_bytes))

            project = meeting.project
            companies = ", ".join([c.name for c in project.companies.all()])
            
            replacements = {
                "{Empresa}": companies,
                "{fechaReunion}": meeting.start_date.strftime("%d/%m/%Y") if meeting.start_date else "",
                "{centro}": project.work_center.name if project.work_center else "",
                "{objetoContrato}": project.description if project.description else "",
                "{fechaInicio}": project.start_date.strftime("%d/%m/%Y") if project.start_date else "",
                "{responsable}": f"{project.contract_manager.first_name} {project.contract_manager.last_name or ''}".strip() if project.contract_manager else "",
                "{contactoPrincipal}": f"{project.main_contact.first_name} {project.main_contact.last_name or ''}".strip() if project.main_contact else "",
                "{tipoCentro}": project.work_center.type if project.work_center else "",
            }

            replace_in_blocks(doc, replacements)

            for section in doc.sections:
                replace_in_blocks(section.header, replacements)
                replace_in_blocks(section.footer, replacements)

            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_base64 = base64.b64encode(output_stream.getvalue()).decode('utf-8')

            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            meeting.document_data = f"data:{mime_type};base64,{output_base64}"
            meeting.save(update_fields=['document_data'])

    @action(detail=True, methods=['post'])
    def notify(self, request, pk=None):
        meeting = self.get_object()
        contacts = list(meeting.notification_contacts.all())

        if not contacts:
            if meeting.project.main_contact:
                contacts = [meeting.project.main_contact]
            else:
                return Response({"error": "No hay destinatarios configurados."}, status=status.HTTP_400_BAD_REQUEST)

        emails = [contact.email for contact in contacts if contact.email]
        if not emails:
            return Response({"error": "Los contactos seleccionados no tienen correo electrónico."}, status=status.HTTP_400_BAD_REQUEST)

        subject = f"Convocatoria de Reunión: {meeting.reason}"
        message = (
            f"Se ha programado una nueva reunión para el proyecto: {meeting.project.description}\n\n"
            f"Motivo: {meeting.reason}\n"
            f"Fecha de inicio: {meeting.start_date}\n"
            f"Hora: {meeting.time}\n"
            f"Tipo: {meeting.type}\n"
            f"Lugar/Enlace: {meeting.teams_link if meeting.type == 'ONLINE' else meeting.location}\n\n"
            f"Por favor, revisa la plataforma para más detalles.\n"
        )

        try:
            send_mail(
                subject=subject,
                message=message,
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=emails,
                fail_silently=False,
            )
            return Response({"message": "Notificación enviada con éxito."}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class DocumentTemplateViewSet(viewsets.ModelViewSet):
    queryset = DocumentTemplate.objects.all()
    serializer_class = DocumentTemplateSerializer

class MeetingDocumentViewSet(viewsets.ModelViewSet):
    queryset = MeetingDocument.objects.all()
    serializer_class = MeetingDocumentSerializer

class FollowUpMeetingViewSet(viewsets.ModelViewSet):
    queryset = FollowUpMeeting.objects.all().order_by('-created_at')
    serializer_class = FollowUpMeetingSerializer

    @action(detail=True, methods=['post'])
    def generate_acta(self, request, pk=None):
        meeting = self.get_object()

        template = DocumentTemplate.objects.filter(category__icontains="seguimiento").first()
        if not template or not template.file_data:
            return Response(
                {"error": "No se encontró la plantilla 'Acta Seguimiento'. Suba la plantilla primero."},
                status=status.HTTP_404_NOT_FOUND
            )

        file_data = template.file_data
        if "," in file_data:
            base64_data = file_data.split(",")[1]
        else:
            base64_data = file_data

        try:
            doc_bytes = base64.b64decode(base64_data)
            doc = Document(io.BytesIO(doc_bytes))
        except Exception as e:
            return Response({"error": f"Error decodificando plantilla: {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        config = getattr(meeting, 'config', None)

        centros = ", ".join([wc.name for wc in meeting.work_centers.all()])
        provincias = ", ".join(meeting.provinces) if meeting.provinces else ""

        replacements = {
            "{fechaReunion}": meeting.date.strftime("%d/%m/%Y") if meeting.date else "",
            "{horaReunión}": meeting.time.strftime("%H:%M") if meeting.time else "",
            "{horaReunion}": meeting.time.strftime("%H:%M") if meeting.time else "",
            "{revision}": config.revision_informacion or "" if config else "",
            "{observacion}": config.observaciones_intercambio or "" if config else "",
            "{concurrencia}": config.solapes_empresas or "" if config else "",
            "{accidentes}": config.accidentes_trabajo or "" if config else "",
            "{emergencia}": config.emergencia or "" if config else "",
            "{temas}": config.otros_temas or "" if config else "",
            "{preguntas}": config.ruegos_preguntas or "" if config else "",
            "{numeroReunion}": config.numero_reunion if config else "",
            "{anoReunion}": meeting.date.year if meeting.date else "",
            "{lugar}": meeting.location or "",
            "{centros}": centros,
            "{provincias}": provincias,
        }

        replace_in_blocks(doc, replacements)

        for section in doc.sections:
            replace_in_blocks(section.header, replacements)
            replace_in_blocks(section.footer, replacements)

        # Process {tablaFirmas}
        found_p = None
        for p in doc.paragraphs:
            if "{tablaFirmas}" in p.text:
                found_p = p
                break

        if found_p:
            signatures = config.signatures if config and hasattr(config, 'signatures') and config.signatures else []
            table = doc.add_table(rows=len(signatures) + 1, cols=4)
            table.style = 'Table Grid'
            table.autofit = False
            
            # Set column widths
            widths = [Inches(2.1), Inches(1.4), Inches(1.4), Inches(1.2)]
            for i, width in enumerate(widths):
                for cell in table.columns[i].cells:
                    cell.width = width

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Empresa"
            hdr_cells[1].text = "Nombre"
            hdr_cells[2].text = "Cargo"
            hdr_cells[3].text = "Firma"
            
            for i, sig in enumerate(signatures):
                row_cells = table.rows[i + 1].cells
                row_cells[0].text = sig.get('company', '')
                row_cells[1].text = sig.get('name', '')
                row_cells[2].text = sig.get('role', '')
                
                sig_data = sig.get('data', '')
                if sig_data and "," in sig_data:
                    try:
                        img_base64 = sig_data.split(",")[1]
                        img_bytes = base64.b64decode(img_base64)
                        img_stream = io.BytesIO(img_bytes)
                        r = row_cells[3].paragraphs[0].add_run()
                        r.add_picture(img_stream, width=Inches(1.2))
                    except Exception as e:
                        print(f"Error inserting signature image: {e}")
            
            found_p._p.addnext(table._tbl)
            found_p.clear()

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_base64 = base64.b64encode(output_stream.getvalue()).decode('utf-8')

        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        meeting.document_data = f"data:{mime_type};base64,{output_base64}"
        meeting.save(update_fields=['document_data'])

        serializer = self.get_serializer(meeting)
        return Response(serializer.data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def notify(self, request, pk=None):
        meeting = self.get_object()
        contacts = list(meeting.notification_contacts.all())

        if not contacts:
            return Response(
                {"error": "No hay destinatarios configurados. Configure los contactos de notificación primero."},
                status=status.HTTP_400_BAD_REQUEST
            )

        emails = [contact.email for contact in contacts if contact.email]
        if not emails:
            return Response(
                {"error": "Los contactos seleccionados no tienen correo electrónico."},
                status=status.HTTP_400_BAD_REQUEST
            )

        contract = meeting.contract
        subject = f"Convocatoria de Reunión de Seguimiento: {meeting.reason}"
        message = (
            f"Se ha programado una reunión de seguimiento para el contrato: {contract.code} — {contract.description}\n\n"
            f"Motivo: {meeting.reason}\n"
            f"Fecha: {meeting.date.strftime('%d/%m/%Y') if meeting.date else ''}\n"
            f"Hora: {meeting.time.strftime('%H:%M') if meeting.time else ''}\n"
            f"Tipo: {meeting.type}\n"
            f"Lugar/Enlace: {meeting.teams_link if meeting.type == 'ONLINE' else (meeting.location or 'No especificado')}\n\n"
            f"Por favor, revisa la plataforma para más detalles.\n"
        )

        try:
            send_mail(
                subject=subject,
                message=message,
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=emails,
                fail_silently=False,
            )
            meeting.is_notified = True
            meeting.save(update_fields=['is_notified'])
            return Response({"message": "Notificación enviada con éxito."}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class FollowUpMeetingConfigViewSet(viewsets.ModelViewSet):
    queryset = FollowUpMeetingConfig.objects.all()
    serializer_class = FollowUpMeetingConfigSerializer

    def create(self, request, *args, **kwargs):
        meeting_id = request.data.get('meeting')
        if meeting_id:
            # Check if a config already exists for this meeting
            existing = FollowUpMeetingConfig.objects.filter(meeting_id=meeting_id).first()
            if existing:
                # If it exists, update it instead of creating a new one
                serializer = self.get_serializer(existing, data=request.data, partial=True)
                serializer.is_valid(raise_exception=True)
                serializer.save()
                return Response(serializer.data, status=status.HTTP_200_OK)
        
        return super().create(request, *args, **kwargs)
